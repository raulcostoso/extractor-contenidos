import streamlit as st
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

# --- Funciones de conversi√≥n HTML a DOCX (sin cambios) ---
def add_styled_run(paragraph_or_heading, bs_node):
    if isinstance(bs_node, NavigableString):
        text = str(bs_node)
        if text.strip():
            paragraph_or_heading.add_run(text)
    elif isinstance(bs_node, Tag):
        text = bs_node.get_text(strip=True)
        if not text and bs_node.name != 'br':
            return

        if bs_node.name in ['strong', 'b']:
            run = paragraph_or_heading.add_run(text)
            run.bold = True
        elif bs_node.name in ['em', 'i']:
            run = paragraph_or_heading.add_run(text)
            run.italic = True
        elif bs_node.name == 'a':
            href = bs_node.get('href')
            if href:
                try:
                    paragraph_or_heading.add_hyperlink(text, href, is_external=True)
                except Exception as e:
                    st.warning(f"Advertencia: No se pudo crear el hiperv√≠nculo para '{text}': {e}")
                    run = paragraph_or_heading.add_run(text + f" [{href}]")
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                    run.font.underline = True
            else:
                paragraph_or_heading.add_run(text)
        elif bs_node.name == 'br':
            paragraph_or_heading.add_run().add_break()
        else:
            if bs_node.contents:
                for child in bs_node.contents:
                    add_styled_run(paragraph_or_heading, child)
            elif text:
                 paragraph_or_heading.add_run(text)

def html_to_docx_elements(bs_element, document_or_container):
    if isinstance(bs_element, NavigableString):
        text = str(bs_element).strip()
        if text:
            p = document_or_container.add_paragraph()
            p.add_run(text)
        return

    if not isinstance(bs_element, Tag):
        return

    tag_name = bs_element.name.lower()

    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag_name[1])
        heading = document_or_container.add_heading(level=level)
        for content_node in bs_element.contents:
            add_styled_run(heading, content_node)
    elif tag_name == 'p':
        p = document_or_container.add_paragraph()
        for content_node in bs_element.contents:
            add_styled_run(p, content_node)
    elif tag_name == 'ul':
        for li in bs_element.find_all('li', recursive=False):
            item_p = document_or_container.add_paragraph(style='ListBullet')
            for content_node in li.contents:
                add_styled_run(item_p, content_node)
    elif tag_name == 'ol':
        for li in bs_element.find_all('li', recursive=False):
            item_p = document_or_container.add_paragraph(style='ListNumber')
            for content_node in li.contents:
                add_styled_run(item_p, content_node)
    elif tag_name == 'br':
        if document_or_container.paragraphs:
            document_or_container.paragraphs[-1].add_run().add_break()
        else:
            document_or_container.add_paragraph().add_run().add_break()
    elif tag_name == 'div':
        for child in bs_element.children:
            html_to_docx_elements(child, document_or_container)
    elif tag_name == 'style':
        pass
    else:
        for child in bs_element.children:
            html_to_docx_elements(child, document_or_container)
# --- Aplicaci√≥n Streamlit ---
st.set_page_config(page_title="HTML a Word", layout="wide")
st.title("üìÑ Extractor de Contenido HTML a Documento Word")
st.markdown("""
Esta aplicaci√≥n te permite extraer el contenido de un elemento HTML espec√≠fico
de una p√°gina web y guardarlo como un documento de Word (.docx).
Puedes especificar el elemento por su ID o por su(s) clase(s).
""")

st.markdown("---")

st.subheader("‚öôÔ∏è Configuraci√≥n de Extracci√≥n")

# Mover el st.radio FUERA del formulario
selection_method = st.radio(
    "M√©todo de selecci√≥n del elemento:",
    ('ID', 'Clase(s) CSS'),
    horizontal=True,
    key="selection_method_radio" # A√±adir una key √∫nica es buena pr√°ctica
)

with st.form(key="extraction_form"):
    url = st.text_input("üîó URL de la p√°gina:", "https://www.unir.net/educacion/master-secundaria/")

    # Los campos de texto ahora reaccionar√°n al radio button
    if selection_method == 'ID':
        target_identifier_value = st.text_input("üÜî ID del div a extraer:", "main-description", key="target_id_input")
        target_identifier_type = "ID"
    elif selection_method == 'Clase(s) CSS':
        target_identifier_value = st.text_input(
            "üè∑Ô∏è Clase(s) del div a extraer (ej: `content main` o `mi-clase-unica`):",
            "list--icons list--numbers list--square list--links magento -margin-bottom--element",
            key="target_class_input"
        )
        st.caption("Si son m√∫ltiples clases, sep√°ralas por espacio. No incluyas el punto `.` inicial.")
        target_identifier_type = "Clase(s)"
    else: # Fallback por si acaso, aunque no deber√≠a ocurrir con radio
        target_identifier_value = ""
        target_identifier_type = "Desconocido"


    submitted = st.form_submit_button("üöÄ Extraer y Convertir")

if submitted: # La l√≥gica de procesamiento sigue igual
    if not url:
        st.error("Por favor, introduce una URL.")
    elif not target_identifier_value:
        # Asegurarse de que target_identifier_type est√° definido incluso si el input no se mostr√≥ (no deber√≠a pasar con este arreglo)
        current_type = "ID" if selection_method == "ID" else "Clase(s)"
        st.error(f"Por favor, introduce un valor para {current_type}.")
    else:
        try:
            with st.spinner(f"Descargando contenido de {url}..."):
                response = requests.get(url, timeout=15)
                response.raise_for_status()
            st.success(f"P√°gina descargada exitosamente de {url}")

            with st.spinner("Parseando HTML..."):
                soup = BeautifulSoup(response.text, 'html.parser')
                main_content_div = None
                actual_search_method = selection_method # Usar el valor del radio que est√° fuera del form

                if actual_search_method == 'ID':
                    main_content_div = soup.find('div', id=target_identifier_value.strip())
                    search_criteria_display = f"ID='{target_identifier_value.strip()}'"
                elif actual_search_method == 'Clase(s) CSS':
                    class_value = target_identifier_value.strip()
                    main_content_div = soup.find('div', class_=class_value)
                    search_criteria_display = f"Clase(s)='{class_value}'"

            if main_content_div:
                st.success(f"Div con {search_criteria_display} encontrado.")

                with st.spinner("Convirtiendo HTML a DOCX..."):
                    document = Document()
                    document.core_properties.title = f"Contenido de {search_criteria_display} de {url}"
                    document.core_properties.author = "Extractor HTML Streamlit App"

                    for element in main_content_div.children:
                        html_to_docx_elements(element, document)

                doc_io = BytesIO()
                document.save(doc_io)
                doc_io.seek(0)

                st.success("¬°Conversi√≥n a Word completada!")

                clean_url_for_filename = url.split('//')[-1].split('/')[0].replace('.', '_')
                clean_identifier = "".join(c if c.isalnum() else "_" for c in target_identifier_value.strip())[:30]
                output_filename = f"contenido_{clean_url_for_filename}_{clean_identifier}.docx"

                st.download_button(
                    label="üì• Descargar Documento Word (.docx)",
                    data=doc_io,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(f"No se encontr√≥ ning√∫n div con {search_criteria_display} en la p√°gina.")
                st.info("Verifica la URL y el identificador. Puedes inspeccionar el c√≥digo fuente de la p√°gina (Ctrl+U o Cmd+Opt+U) para encontrar los valores correctos.")

        except requests.exceptions.RequestException as e:
            st.error(f"Error de red al intentar acceder a la URL: {e}")
        except Exception as e:
            st.error(f"Ocurri√≥ un error inesperado durante el proceso.")
            st.exception(e)

st.sidebar.markdown("---")
st.sidebar.info("Creado con Streamlit y python-docx.")import streamlit as st
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

# --- Funciones de conversi√≥n HTML a DOCX (sin cambios) ---
def add_styled_run(paragraph_or_heading, bs_node):
    if isinstance(bs_node, NavigableString):
        text = str(bs_node)
        if text.strip():
            paragraph_or_heading.add_run(text)
    elif isinstance(bs_node, Tag):
        text = bs_node.get_text(strip=True)
        if not text and bs_node.name != 'br':
            return

        if bs_node.name in ['strong', 'b']:
            run = paragraph_or_heading.add_run(text)
            run.bold = True
        elif bs_node.name in ['em', 'i']:
            run = paragraph_or_heading.add_run(text)
            run.italic = True
        elif bs_node.name == 'a':
            href = bs_node.get('href')
            if href:
                try:
                    paragraph_or_heading.add_hyperlink(text, href, is_external=True)
                except Exception as e:
                    st.warning(f"Advertencia: No se pudo crear el hiperv√≠nculo para '{text}': {e}")
                    run = paragraph_or_heading.add_run(text + f" [{href}]")
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                    run.font.underline = True
            else:
                paragraph_or_heading.add_run(text)
        elif bs_node.name == 'br':
            paragraph_or_heading.add_run().add_break()
        else:
            if bs_node.contents:
                for child in bs_node.contents:
                    add_styled_run(paragraph_or_heading, child)
            elif text:
                 paragraph_or_heading.add_run(text)

def html_to_docx_elements(bs_element, document_or_container):
    if isinstance(bs_element, NavigableString):
        text = str(bs_element).strip()
        if text:
            p = document_or_container.add_paragraph()
            p.add_run(text)
        return

    if not isinstance(bs_element, Tag):
        return

    tag_name = bs_element.name.lower()

    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag_name[1])
        heading = document_or_container.add_heading(level=level)
        for content_node in bs_element.contents:
            add_styled_run(heading, content_node)
    elif tag_name == 'p':
        p = document_or_container.add_paragraph()
        for content_node in bs_element.contents:
            add_styled_run(p, content_node)
    elif tag_name == 'ul':
        for li in bs_element.find_all('li', recursive=False):
            item_p = document_or_container.add_paragraph(style='ListBullet')
            for content_node in li.contents:
                add_styled_run(item_p, content_node)
    elif tag_name == 'ol':
        for li in bs_element.find_all('li', recursive=False):
            item_p = document_or_container.add_paragraph(style='ListNumber')
            for content_node in li.contents:
                add_styled_run(item_p, content_node)
    elif tag_name == 'br':
        if document_or_container.paragraphs:
            document_or_container.paragraphs[-1].add_run().add_break()
        else:
            document_or_container.add_paragraph().add_run().add_break()
    elif tag_name == 'div':
        for child in bs_element.children:
            html_to_docx_elements(child, document_or_container)
    elif tag_name == 'style':
        pass
    else:
        for child in bs_element.children:
            html_to_docx_elements(child, document_or_container)
# --- Aplicaci√≥n Streamlit ---
st.set_page_config(page_title="HTML a Word", layout="wide")
st.title("üìÑ Extractor de Contenido HTML a Documento Word")
st.markdown("""
Esta aplicaci√≥n te permite extraer el contenido de un elemento HTML espec√≠fico
de una p√°gina web y guardarlo como un documento de Word (.docx).
Puedes especificar el elemento por su ID o por su(s) clase(s).
""")

st.markdown("---")

st.subheader("‚öôÔ∏è Configuraci√≥n de Extracci√≥n")

# Mover el st.radio FUERA del formulario
selection_method = st.radio(
    "M√©todo de selecci√≥n del elemento:",
    ('ID', 'Clase(s) CSS'),
    horizontal=True,
    key="selection_method_radio" # A√±adir una key √∫nica es buena pr√°ctica
)

with st.form(key="extraction_form"):
    url = st.text_input("üîó URL de la p√°gina:", "https://www.unir.net/educacion/master-secundaria/")

    # Los campos de texto ahora reaccionar√°n al radio button
    if selection_method == 'ID':
        target_identifier_value = st.text_input("üÜî ID del div a extraer:", "main-description", key="target_id_input")
        target_identifier_type = "ID"
    elif selection_method == 'Clase(s) CSS':
        target_identifier_value = st.text_input(
            "üè∑Ô∏è Clase(s) del div a extraer (ej: `content main` o `mi-clase-unica`):",
            "list--icons list--numbers list--square list--links magento -margin-bottom--element",
            key="target_class_input"
        )
        st.caption("Si son m√∫ltiples clases, sep√°ralas por espacio. No incluyas el punto `.` inicial.")
        target_identifier_type = "Clase(s)"
    else: # Fallback por si acaso, aunque no deber√≠a ocurrir con radio
        target_identifier_value = ""
        target_identifier_type = "Desconocido"


    submitted = st.form_submit_button("üöÄ Extraer y Convertir")

if submitted: # La l√≥gica de procesamiento sigue igual
    if not url:
        st.error("Por favor, introduce una URL.")
    elif not target_identifier_value:
        # Asegurarse de que target_identifier_type est√° definido incluso si el input no se mostr√≥ (no deber√≠a pasar con este arreglo)
        current_type = "ID" if selection_method == "ID" else "Clase(s)"
        st.error(f"Por favor, introduce un valor para {current_type}.")
    else:
        try:
            with st.spinner(f"Descargando contenido de {url}..."):
                response = requests.get(url, timeout=15)
                response.raise_for_status()
            st.success(f"P√°gina descargada exitosamente de {url}")

            with st.spinner("Parseando HTML..."):
                soup = BeautifulSoup(response.text, 'html.parser')
                main_content_div = None
                actual_search_method = selection_method # Usar el valor del radio que est√° fuera del form

                if actual_search_method == 'ID':
                    main_content_div = soup.find('div', id=target_identifier_value.strip())
                    search_criteria_display = f"ID='{target_identifier_value.strip()}'"
                elif actual_search_method == 'Clase(s) CSS':
                    class_value = target_identifier_value.strip()
                    main_content_div = soup.find('div', class_=class_value)
                    search_criteria_display = f"Clase(s)='{class_value}'"

            if main_content_div:
                st.success(f"Div con {search_criteria_display} encontrado.")

                with st.spinner("Convirtiendo HTML a DOCX..."):
                    document = Document()
                    document.core_properties.title = f"Contenido de {search_criteria_display} de {url}"
                    document.core_properties.author = "Extractor HTML Streamlit App"

                    for element in main_content_div.children:
                        html_to_docx_elements(element, document)

                doc_io = BytesIO()
                document.save(doc_io)
                doc_io.seek(0)

                st.success("¬°Conversi√≥n a Word completada!")

                clean_url_for_filename = url.split('//')[-1].split('/')[0].replace('.', '_')
                clean_identifier = "".join(c if c.isalnum() else "_" for c in target_identifier_value.strip())[:30]
                output_filename = f"contenido_{clean_url_for_filename}_{clean_identifier}.docx"

                st.download_button(
                    label="üì• Descargar Documento Word (.docx)",
                    data=doc_io,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(f"No se encontr√≥ ning√∫n div con {search_criteria_display} en la p√°gina.")
                st.info("Verifica la URL y el identificador. Puedes inspeccionar el c√≥digo fuente de la p√°gina (Ctrl+U o Cmd+Opt+U) para encontrar los valores correctos.")

        except requests.exceptions.RequestException as e:
            st.error(f"Error de red al intentar acceder a la URL: {e}")
        except Exception as e:
            st.error(f"Ocurri√≥ un error inesperado durante el proceso.")
            st.exception(e)

st.sidebar.markdown("---")
st.sidebar.info("Creado con Streamlit y python-docx.")
